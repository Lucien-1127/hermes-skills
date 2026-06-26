#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generates a Taiwan legal document (暫緩執行聲請書) in Word (.docx) format.
Usage:  python3 這支腳本.py
Output: ~/暫緩執行聲請書_115執6265_丑股_專業版.docx

To convert to PDF:
  libreoffice --headless --convert-to pdf ~/暫緩執行聲請書_115執6265_丑股_專業版.docx

Requirements: pip3 install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── page setup ──
s = doc.sections[0]
s.page_width = Cm(21.0); s.page_height = Cm(29.7)
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

def set_font(run, size=14, bold=False):
    run.font.name = 'DFKai-SB'; run.font.size = Pt(size); run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), 'DFKai-SB')

def P(text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, sa=0, indent=None):
    para = doc.add_paragraph(); para.alignment = align
    pf = para.paragraph_format; pf.space_after = Pt(sa); pf.line_spacing = Pt(28)
    if indent: pf.first_line_indent = Cm(indent)
    run = para.add_run(text); set_font(run, size, bold); return para

def page_break(): doc.add_page_break()

def add_page_number():
    footer = s.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); set_font(run, size=10)
    for t in ['begin', 'end']:
        fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), t)
        run._element.append(fld)
        if t == 'begin':
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
            run._element.append(instr)

# ── title ──
P('暫緩執行聲請書', size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=18)
P('案　　號：115 年度執字第 6265 號', sa=2)
P('股　　別：丑股', sa=10)
P('聲 請 人（即受刑人）：楊瑞忠', sa=2)
P('出生日期：民國 68 年 5 月 14 日', sa=2)
P('國民身分證統一編號：＿＿＿＿＿＿＿＿＿＿（請填寫）', sa=2)
P('聯絡電話：＿＿＿＿＿＿＿＿＿＿（請填寫）', sa=2)
P('送達住址：桃園市大園區中正東路一段 739 號', sa=14)
P('為聲請准予暫緩（延緩）執行事：', bold=True, sa=10)

# ── body ──
P('一、執行命令摘要', bold=True, sa=6)
P('聲請人前因竊盜案件，經鈞署以 115 年度執字第 6265 號執行傳票通知，應於民國 115 年 7 月 14 日上午 9 時報到執行。聲請人前因竊盜罪，經判處有期徒刑＿＿＿＿＿月確定，現等候執行在案。', sa=10, indent=0.8)

P('二、暫緩執行之事由——母親乏人照料之急迫困境', bold=True, sa=6)
P('聲請人之母親○○○（民國 44 年＿＿月＿＿日生）年事已高，患有＿＿＿＿＿＿＿＿＿＿＿＿（病名詳如附件二診斷證明書所載），生活已無法自理，須仰賴聲請人全天候照料。', sa=4, indent=0.8)
P('聲請人為家中唯一實際照護母親之人。聲請人之兄長長年旅居國外（旅外證明如附件五），未曾返臺協助照料；弟弟○○○現從事＿＿＿＿＿＿＿＿工作，每日工作時間自＿＿時至＿＿時，且育有未成年子女＿＿名，自身家庭照顧負擔沉重，實際上無力兼顧母親全天候之生活照護（弟弟在職證明如附件六、戶籍謄本如附件三）。聲請人現與母親同住於租賃處所（租賃契約如附件四），倘聲請人於此時逕行入監，母親頓失依靠，將陷於無人照顧之急迫困境。', sa=10, indent=0.8)

P('三、法律依據', bold=True, sa=6)
P('按刑事訴訟法第 457 條第 1 項規定，執行裁判由為裁判法院對應之檢察官指揮之；檢察官就執行指揮之具體時程，本有依個案情形裁量酌定之空間。受刑人如有重大家庭變故，致不能即時到案執行者，檢察官非不得本於職權，酌予延緩執行之日期。聲請人所陳情形，核屬上開規定所欲斟酌之範疇。', sa=4, indent=0.8)
P('末按憲法第 23 條所揭比例原則，國家公權力之行使，於可兼顧之範圍內，應選擇侵害最小之手段。聲請人並非規避執行，僅因母親照護一事亟待安排，懇請鈞署於不影響執行終局之前提下，惠予短期延緩，俾免聲請人之母親因執行而頓陷乏人照料之危殆。', sa=10, indent=0.8)

P('四、具體請求', bold=True, sa=6)
P('聲請人刻正積極為母親辦理安養機構入住事宜，業已與＿＿＿＿＿＿＿＿＿＿安養機構接洽，預訂於民國 115 年＿＿月＿＿日辦理入住手續（接洽紀錄如附件七）。按安養機構一般收住程序，須歷經健康評估、候位及契約簽訂等流程，通常約需 60 日。為使母親於聲請人執行期間獲得妥適照顧，懇請鈞署惠予將執行日期延展至民國 115 年 8 月 25 日。', sa=4, indent=0.8)
P('聲請人謹此承諾，俟母親安置妥當，必如期報到執行，勇於面對法律責任，絕不藉故拖延或逃避。', sa=14, indent=0.8)

P('此　　致', sa=6)
P('臺灣桃園地方檢察署　執行科　公鑒', sa=16)

P('附件：', bold=True, sa=6)
for i, a in enumerate([
    '一、執行傳票影本一份。',
    '二、母親之診斷證明書（近期開立，載明病名及須專人照顧）一份。',
    '三、戶籍謄本（載明家屬狀況及弟弟子女數）一份。',
    '四、租賃契約影本（證明聲請人與母親同住）一份。',
    '五、兄長旅居國外之證明文件一份。',
    '六、弟弟在職證明（證明工時無法全日照護）一份。',
    '七、安養機構接洽紀錄或訂金單據影本一份。',
], 1):
    P(a, sa=2, indent=0.8)

P('', sa=14)
P('聲　請　人：楊瑞忠　（簽名蓋章）', sa=4)
P('送達代收人：＿＿＿＿＿＿＿＿＿＿', sa=4)
P('聯絡電話：＿＿＿＿＿＿＿＿＿＿', sa=4)
P('', sa=4)
P('中　華　民　國　115　年　＿＿　月　＿＿　日', align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

# ── page 2: instructions ──
page_break()
P('【填寫說明】', size=16, bold=True, sa=10)
for n in [
    '一、畫底線「＿＿」處請依實際狀況填寫，務必與證明文件一致。',
    '二、刑期月數：請填寫判決書主文所載之有期徒刑月數。',
    '三、母親姓名及出生月日：應與戶籍謄本及診斷證明書逐字相符。',
    '四、病名：應與診斷證明書所載診斷名稱完全相同。',
    '五、弟弟姓名、工作、工時、子女數：依實際狀況填寫，並備在職證明。',
    '六、安養機構名稱：已接洽之機構全名，入住日依機構安排填寫。',
    '七、身分證字號、電話、送達代收人：送件前務必填妥。',
    '八、簽名蓋章：聲請人欄須親簽並蓋章。',
]:
    P(n, sa=3, indent=0.8)

P('', sa=8)
P('【送件提醒】', size=16, bold=True, sa=6)
for r in [
    '一、診斷證明書應於 1 個月內開立，載明「需專人照顧、不宜獨居」等醫囑。',
    '二、附件齊全後，以掛號寄送或親送至桃園地檢署服務台。',
    '三、送件後 2-3 日請電 1096 分機確認收文及承辦股。',
    '四、原報到日（7 月 14 日）在未獲准通知前仍有效，切勿擅自不到。',
    '五、本文件為書狀草稿，不構成正式法律意見。',
]:
    P(r, sa=3, indent=0.8)

add_page_number()

out = os.path.expanduser('~/暫緩執行聲請書_115執6265_丑股_專業版.docx')
doc.save(out)
print(f'✅ Saved: {out}')
